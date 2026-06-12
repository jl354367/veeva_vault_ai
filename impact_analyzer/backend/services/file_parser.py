"""Unified file parser — handles xlsx, xls, csv, pdf, docx, txt, json."""
import io
import json
import traceback
import pandas as pd

ALLOWED_EXTENSIONS = {".xlsx", ".xls", ".csv", ".pdf", ".docx", ".txt", ".json"}


def parse_file(content: bytes, ext: str) -> tuple[str, dict]:
    """Return (full_text, {sheet_name: DataFrame}) for any supported format."""
    ext = ext.lower()
    if ext in (".xlsx", ".xls"):
        return _parse_excel(content)
    if ext == ".csv":
        return _parse_csv(content)
    if ext == ".pdf":
        return _parse_pdf(content)
    if ext == ".docx":
        return _parse_docx(content)
    if ext == ".txt":
        return _parse_txt(content)
    if ext == ".json":
        return _parse_json(content)
    raise ValueError(f"Unsupported file type: {ext}")


def _parse_excel(content: bytes) -> tuple[str, dict]:
    xls = pd.ExcelFile(io.BytesIO(content))
    sheets, parts = {}, []
    for sheet in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet, engine="openpyxl")
        df.dropna(how="all", inplace=True)
        df.dropna(axis=1, how="all", inplace=True)
        if df.empty:
            continue
        sheets[sheet] = df
        parts.append(f"=== Sheet: {sheet} ===\n{df.to_string(index=False)}")
    return "\n\n".join(parts) if parts else "(No data found in file)", sheets


def _parse_csv(content: bytes) -> tuple[str, dict]:
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            df = pd.read_csv(io.BytesIO(content), encoding=enc)
            break
        except Exception:
            continue
    else:
        df = pd.read_csv(io.BytesIO(content), encoding="utf-8", errors="replace")
    df.dropna(how="all", inplace=True)
    df.dropna(axis=1, how="all", inplace=True)
    return df.to_string(index=False), {"Sheet1": df}


def _parse_pdf(content: bytes) -> tuple[str, dict]:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("Install pdfplumber: pip install pdfplumber")

    sheets, text_parts = {}, []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            page_text = page.extract_text() or ""
            tables = page.extract_tables() or []

            for j, table in enumerate(tables, 1):
                if not table or len(table) < 2:
                    continue
                headers = [str(c) if c is not None else f"Col{k}" for k, c in enumerate(table[0])]
                rows = [[str(c) if c is not None else "" for c in row] for row in table[1:]]
                df = pd.DataFrame(rows, columns=headers)
                df.dropna(how="all", inplace=True)
                df.dropna(axis=1, how="all", inplace=True)
                if df.empty:
                    continue
                key = f"Page{i}" if len(tables) == 1 else f"Page{i}_T{j}"
                sheets[key] = df
                text_parts.append(f"=== {key} ===\n{df.to_string(index=False)}")

            if page_text.strip():
                text_parts.append(f"=== Page {i} ===\n{page_text}")

    if not sheets:
        full_text = "\n\n".join(text_parts) or "(No readable content in PDF)"
        lines = [l for l in full_text.splitlines() if l.strip()]
        sheets["Document"] = pd.DataFrame({"Content": lines})

    return "\n\n".join(text_parts) or "(No content)", sheets


def _parse_docx(content: bytes) -> tuple[str, dict]:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Install python-docx: pip install python-docx")

    doc = Document(io.BytesIO(content))
    sheets, text_parts = {}, []

    for i, table in enumerate(doc.tables, 1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if len(rows) < 2:
            continue
        seen: dict = {}
        clean_hdrs = []
        for h in rows[0]:
            h = h or f"Col{len(clean_hdrs)}"
            if h in seen:
                seen[h] += 1
                h = f"{h}_{seen[h]}"
            else:
                seen[h] = 0
            clean_hdrs.append(h)
        df = pd.DataFrame(rows[1:], columns=clean_hdrs)
        df.dropna(how="all", inplace=True)
        df.dropna(axis=1, how="all", inplace=True)
        if df.empty:
            continue
        sheets[f"Table{i}"] = df
        text_parts.append(f"=== Table {i} ===\n{df.to_string(index=False)}")

    para_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if para_lines:
        text_parts.append("=== Document Text ===\n" + "\n".join(para_lines))
        if not sheets:
            sheets["Document"] = pd.DataFrame({"Content": para_lines})

    return "\n\n".join(text_parts) or "(No content)", sheets


def _parse_txt(content: bytes) -> tuple[str, dict]:
    text = content.decode("utf-8", errors="replace")
    lines = [l for l in text.splitlines() if l.strip()]
    return text, {"Document": pd.DataFrame({"Content": lines})}


def _parse_json(content: bytes) -> tuple[str, dict]:
    data = json.loads(content.decode("utf-8"))
    sheets, parts = {}, []

    if isinstance(data, list):
        df = pd.DataFrame(data)
        df.dropna(how="all", inplace=True)
        return df.to_string(index=False), {"Sheet1": df}

    if isinstance(data, dict):
        for key, val in data.items():
            if isinstance(val, list) and val:
                try:
                    df = pd.DataFrame(val)
                    df.dropna(how="all", inplace=True)
                    sheets[str(key)] = df
                    parts.append(f"=== {key} ===\n{df.to_string(index=False)}")
                except Exception:
                    continue
        if not sheets:
            df = pd.DataFrame([data])
            sheets["Sheet1"] = df
            parts.append(df.to_string(index=False))

    return "\n\n".join(parts) or str(data), sheets
