"""
Parse uploaded documents (PDF, DOCX, XLSX, TXT) into plain text.
"""

from __future__ import annotations

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text(file_path: str | Path) -> str:
    """
    Auto-detect file type and extract all text content.
    Supports: .pdf, .docx, .doc, .xlsx, .xls, .txt, .md
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(path)
    elif suffix in (".docx", ".doc"):
        return _extract_docx(path)
    elif suffix in (".xlsx", ".xls"):
        return _extract_excel(path)
    elif suffix in (".txt", ".md", ".text"):
        return _extract_text_file(path)
    else:
        logger.warning("Unknown file type %s — attempting plain-text read", suffix)
        return _extract_text_file(path)


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """Parse from raw bytes (e.g. when the file hasn't been saved yet)."""
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf_bytes(content)
    elif suffix in (".docx", ".doc"):
        return _extract_docx_bytes(content)
    elif suffix in (".xlsx", ".xls"):
        return _extract_excel_bytes(content)
    else:
        return content.decode("utf-8", errors="replace")


# ─── PDF ──────────────────────────────────────────────────────────────────────

def _extract_pdf(path: Path) -> str:
    with open(path, "rb") as f:
        return _extract_pdf_bytes(f.read())


def _extract_pdf_bytes(content: bytes) -> str:
    try:
        import PyPDF2

        reader = PyPDF2.PdfReader(io.BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n\n".join(pages).strip()
        if not text:
            logger.warning("PDF extraction returned empty text — may be image-based")
        return text
    except Exception as exc:
        logger.error("PDF extraction failed: %s", exc)
        raise ValueError(f"Could not extract text from PDF: {exc}") from exc


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def _extract_docx(path: Path) -> str:
    with open(path, "rb") as f:
        return _extract_docx_bytes(f.read())


def _extract_docx_bytes(content: bytes) -> str:
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also grab text from tables
        table_texts: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        all_text = paragraphs + table_texts
        return "\n\n".join(all_text).strip()
    except Exception as exc:
        logger.error("DOCX extraction failed: %s", exc)
        raise ValueError(f"Could not extract text from DOCX: {exc}") from exc


# ─── Excel (XLSX / XLS) ───────────────────────────────────────────────────────

def _extract_excel(path: Path) -> str:
    with open(path, "rb") as f:
        return _extract_excel_bytes(f.read())


def _extract_excel_bytes(content: bytes) -> str:
    try:
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        sections: list[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows: list[str] = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                sections.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows))
        wb.close()
        result = "\n\n".join(sections).strip()
        if not result:
            logger.warning("Excel extraction returned empty text")
        return result
    except Exception as exc:
        logger.error("Excel extraction failed: %s", exc)
        raise ValueError(f"Could not extract text from Excel file: {exc}") from exc


# ─── Plain text ───────────────────────────────────────────────────────────────

def _extract_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace").strip()
    except Exception as exc:
        logger.error("Text file read failed: %s", exc)
        raise ValueError(f"Could not read text file: {exc}") from exc
